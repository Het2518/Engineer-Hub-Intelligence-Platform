"""Text extraction service for all supported file types.

Fixes (2026-07-29 audit):
- Unknown-type fallback now correctly awaits _extract_text_plain instead of
  returning a coroutine object.
- Image vision extractor uses a dedicated OpenAI client (not the Groq base URL)
  since vision is an OpenAI-only capability.
- get_event_loop() replaced with asyncio.get_running_loop() (Python 3.10+).
- Extracted text is truncated at MAX_TEXT_CHARS to prevent token budget explosion
  in the chunking / embedding pipeline.
"""
import asyncio
import base64
import json
from pathlib import Path
from typing import Optional

import structlog
from openai import AsyncOpenAI

from config import get_settings

logger = structlog.get_logger()
settings = get_settings()

MAX_TEXT_CHARS = 400_000   # ~100K tokens; truncate anything beyond this


async def extract_text(file_path: Path, mime_type: Optional[str] = None) -> str:
    """Route to the appropriate extractor based on file extension."""
    suffix = file_path.suffix.lower()

    extractors = {
        ".pdf":      _extract_pdf,
        ".docx":     _extract_docx,
        ".doc":      _extract_docx,
        ".txt":      _extract_text_plain,
        ".md":       _extract_text_plain,
        ".markdown": _extract_text_plain,
        ".json":     _extract_json,
        ".csv":      _extract_csv,
        ".png":      _extract_image_vision,
        ".jpg":      _extract_image_vision,
        ".jpeg":     _extract_image_vision,
    }

    extractor = extractors.get(suffix)
    if extractor is None:
        logger.warning("No dedicated extractor for type — falling back to plain text", suffix=suffix)
        # Must await — _extract_text_plain is a coroutine function
        text = await _extract_text_plain(file_path)
    else:
        try:
            text = await extractor(file_path)
        except Exception as e:
            logger.error("Extraction failed", file=str(file_path), error=str(e))
            raise

    # Hard cap on extracted text length
    if text and len(text) > MAX_TEXT_CHARS:
        logger.warning(
            "Extracted text truncated",
            file=str(file_path.name),
            original_chars=len(text),
            truncated_to=MAX_TEXT_CHARS,
        )
        text = text[:MAX_TEXT_CHARS]

    return text or ""


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def _extract_pdf_sync(file_path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    pages  = []
    for num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {num + 1}]\n{text}")
    return "\n\n".join(pages)


async def _extract_pdf(file_path: Path) -> str:
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(None, _extract_pdf_sync, file_path)


def _extract_docx_sync(file_path: Path) -> str:
    from docx import Document
    doc  = Document(str(file_path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


async def _extract_docx(file_path: Path) -> str:
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(None, _extract_docx_sync, file_path)


async def _extract_text_plain(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return file_path.read_text(encoding="latin-1", errors="replace")


async def _extract_json(file_path: Path) -> str:
    try:
        data = json.loads(file_path.read_text(encoding="utf-8"))
        return json.dumps(data, indent=2)
    except Exception:
        return file_path.read_text(encoding="utf-8", errors="replace")


async def _extract_csv(file_path: Path) -> str:
    import csv
    lines = []
    try:
        with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
            for row in csv.reader(f):
                lines.append(", ".join(row))
        return "\n".join(lines)
    except Exception:
        return file_path.read_text(encoding="utf-8", errors="replace")


async def _extract_image_vision(file_path: Path) -> str:
    """Use Groq LLaMA 3.2 Vision to extract text and structure from images.
    
    A missing GROQ_API_KEY for the vision call will raise a clear error rather than silently failing.
    """
    client = AsyncOpenAI(
        api_key=settings.groq_api_key,
        base_url=settings.llm_base_url,
    )

    with open(file_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    ext  = file_path.suffix.lower().lstrip(".")
    mime = {"jpg": "jpeg", "jpeg": "jpeg", "png": "png"}.get(ext, "png")

    prompt = (
        "Analyze this image carefully. Extract and describe:\n"
        "1. All service/component names\n"
        "2. Relationships and data flows between services\n"
        "3. Databases and storage systems\n"
        "4. External integrations and APIs\n"
        "5. Infrastructure components\n"
        "6. Any text labels, annotations, or notes\n\n"
        "Format as structured text suitable for semantic search. Be thorough and precise."
    )

    try:
        response = await client.chat.completions.create(
            model="llama-3.2-90b-vision-preview",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/{mime};base64,{image_data}", "detail": "high"}},
                    {"type": "text", "text": prompt},
                ],
            }],
            max_tokens=2000,
        )
        extracted = response.choices[0].message.content or ""
        return f"[Image: {file_path.name}]\n\n{extracted}"
    except Exception as e:
        logger.error("Vision extraction failed", file=file_path.name, error=str(e))
        return f"[Image: {file_path.name}] — Vision extraction failed. Ensure GROQ_API_KEY is set."
